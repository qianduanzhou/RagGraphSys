"""Document parsing utilities.

The API layer still consumes a plain string through :func:`parse_upload`, but
the implementation now parses into structured elements first.  Keeping the
intermediate structure makes it easier to preserve page / sheet / table
boundaries and to add stronger parsers later without changing the public API.
"""
from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass, field
from datetime import datetime
from itertools import zip_longest
from typing import Any, Iterable, Set

from core.logger import get_logger

logger = get_logger(__name__)


# ---------------------------------------------------------------------- #
# Supported file types
# ---------------------------------------------------------------------- #
TEXT_EXTS: Set[str] = {
    ".txt", ".md", ".markdown", ".json", ".log", ".rst", ".org",
    ".js", ".jsx", ".ts", ".tsx", ".html", ".htm", ".css", ".scss", ".sass",
    ".less", ".vue", ".svelte",
    ".py", ".java", ".c", ".cpp", ".cc", ".cxx", ".h", ".hpp", ".cs", ".go",
    ".rs", ".kt", ".kts", ".scala", ".swift", ".rb", ".php", ".pl", ".lua",
    ".dart", ".r", ".m", ".mm",
    ".sh", ".bash", ".zsh", ".fish", ".bat", ".cmd", ".ps1",
    ".sql", ".xml", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".properties", ".env", ".gradle", ".gemspec",
}

PDF_EXTS: Set[str] = {".pdf"}
WORD_EXTS: Set[str] = {".docx"}  # legacy .doc needs system-level conversion
CSV_EXTS: Set[str] = {".csv"}
EXCEL_EXTS: Set[str] = {".xlsx", ".xls"}

ALLOWED_EXTS: Set[str] = TEXT_EXTS | CSV_EXTS | PDF_EXTS | WORD_EXTS | EXCEL_EXTS

# Safety limits for large tables.  Values are intentionally kept compatible
# with the old parser, but truncation is now surfaced in parser warnings.
TABLE_MAX_COLS = 50
TABLE_MAX_ROWS = 5000


@dataclass
class ParsedElement:
    """A single parsed document element."""

    kind: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Structured parse result used internally and by tests/tools."""

    text: str
    elements: list[ParsedElement] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _ext(filename: str) -> str:
    name = (filename or "").lower()
    if "." not in name:
        return ""
    return "." + name.rsplit(".", 1)[-1]


def parse_upload(filename: str, raw: bytes) -> str:
    """Parse uploaded file bytes into text for ingestion."""

    parsed = parse_upload_detail(filename, raw)
    text = (parsed.text or "").strip()
    if not text:
        raise ValueError(
            f"未能从 {filename} 提取出任何文本（可能是扫描件 / 图片 / 空文件）"
        )
    return text


def parse_upload_detail(filename: str, raw: bytes) -> ParsedDocument:
    """Parse uploaded bytes and keep lightweight structural metadata."""

    ext = _ext(filename)
    if ext in TEXT_EXTS:
        text = _decode_text_bytes(raw)
        return ParsedDocument(text=text, elements=[ParsedElement("text", text)])
    if ext in CSV_EXTS:
        return _parse_csv_document(raw)
    if ext in PDF_EXTS:
        return _parse_pdf_document(raw)
    if ext in WORD_EXTS:
        return _parse_docx_document(raw)
    if ext in EXCEL_EXTS:
        return _parse_excel_document(ext, raw)

    raise ValueError(
        f"不支持的文件类型 '{ext or '(无扩展名)'}'；"
        f"支持：文本/代码、CSV、PDF、Word(.docx)、Excel(.xlsx/.xls)"
    )


def _decode_text_bytes(raw: bytes) -> str:
    """Decode text-like files with common encodings.

    UTF-8 is preferred, but many Chinese CSV/TXT uploads are GB18030 encoded.
    Falling back before ``errors='ignore'`` avoids silently dropping content.
    """

    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


# ---------------------------------------------------------------------- #
# PDF
# ---------------------------------------------------------------------- #
def _parse_pdf_document(raw: bytes) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # noqa: BLE001
        fallback = _parse_pdf_with_pymupdf(raw)
        if fallback is not None:
            return fallback
        raise ValueError("解析 PDF 需要安装 pypdf：pip install pypdf") from exc

    try:
        reader = PdfReader(io.BytesIO(raw))
        pages: list[tuple[int, str]] = []
        warnings: list[str] = []
        for idx, page in enumerate(reader.pages, 1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append((idx, text))
            else:
                warnings.append(f"PDF 第 {idx} 页未抽取到文本，可能是扫描页或图片页")

        text = _render_page_text(pages)
        expanded_notes = _expand_merged_pdf_table_notes(text)
        if expanded_notes:
            text = f"{text}\n\n{expanded_notes}".strip()

        if text:
            elements = [
                ParsedElement("page", page_text, {"page": page_no})
                for page_no, page_text in pages
            ]
            return ParsedDocument(text=text, elements=elements, warnings=warnings)

        fallback = _parse_pdf_with_pymupdf(raw)
        if fallback is not None:
            fallback.warnings = warnings + fallback.warnings
            return fallback
        return ParsedDocument(text="", warnings=warnings)
    except Exception as exc:  # noqa: BLE001
        logger.warning("PDF parse failed with pypdf: %s", exc)
        fallback = _parse_pdf_with_pymupdf(raw)
        if fallback is not None:
            fallback.warnings.insert(0, f"pypdf 解析失败，已使用 PyMuPDF 兜底：{exc}")
            return fallback
        raise ValueError(f"PDF 解析失败：{exc}") from exc


def _parse_pdf_with_pymupdf(raw: bytes) -> ParsedDocument | None:
    """Optional PyMuPDF fallback.

    PyMuPDF is intentionally optional.  If it is installed in the deployment
    environment, it often handles layout-heavy PDFs better than pypdf.
    """

    try:
        import fitz  # type: ignore[import-not-found]
    except ImportError:
        return None

    try:
        doc = fitz.open(stream=raw, filetype="pdf")
        pages: list[tuple[int, str]] = []
        warnings: list[str] = []
        for idx, page in enumerate(doc, 1):
            text = (page.get_text("text") or "").strip()
            if text:
                pages.append((idx, text))
            else:
                warnings.append(f"PDF 第 {idx} 页未抽取到文本，可能需要 OCR")
        text = _render_page_text(pages)
        expanded_notes = _expand_merged_pdf_table_notes(text)
        if expanded_notes:
            text = f"{text}\n\n{expanded_notes}".strip()
        elements = [
            ParsedElement("page", page_text, {"page": page_no, "engine": "pymupdf"})
            for page_no, page_text in pages
        ]
        return ParsedDocument(text=text, elements=elements, warnings=warnings)
    except Exception as exc:  # noqa: BLE001
        logger.warning("PDF parse failed with PyMuPDF: %s", exc)
        return None


def _render_page_text(pages: list[tuple[int, str]]) -> str:
    if not pages:
        return ""
    if len(pages) == 1:
        return pages[0][1]
    return "\n\n".join(f"## Page {page_no}\n{page_text}" for page_no, page_text in pages)


_JOB_ROW_RE = re.compile(r"(?<!\w)(\d{1,3})\s+(G\d{10})\s+")
_MERGED_NOTE_RE = re.compile(r"适应[^。\n；;]{0,30}(?:[。；;])?")


def _expand_merged_pdf_table_notes(text: str) -> str:
    """Add searchable rows for PDF table notes stored as merged cells."""

    if not text or "适应" not in text:
        return ""

    expanded: list[str] = []
    seen: set[tuple[str, str]] = set()

    for note_match in _MERGED_NOTE_RE.finditer(text):
        note = _normalize_ws(note_match.group(0)).rstrip("；;")
        if not note:
            continue

        block = _subtotal_block_around(text, note_match.start(), note_match.end())
        rows = _extract_job_rows(block)
        if len(rows) <= 1:
            continue

        for row in rows:
            key = (row["code"], note)
            if key in seen:
                continue
            seen.add(key)
            content = _normalize_ws(row["content"].replace(note, " "))
            expanded.append(f'{row["no"]} {row["code"]} {content} | 备注：{note}')

    if not expanded:
        return ""
    return "【PDF表格合并备注展开】\n" + "\n".join(expanded)


def _subtotal_block_around(text: str, start: int, end: int) -> str:
    before = text.rfind("小计", 0, start)
    after = text.find("小计", end)
    block_start = before + len("小计") if before >= 0 else max(0, start - 1200)
    block_end = after if after >= 0 else min(len(text), end + 1200)
    return text[block_start:block_end]


def _extract_job_rows(block: str) -> list[dict[str, str]]:
    matches = list(_JOB_ROW_RE.finditer(block))
    rows: list[dict[str, str]] = []
    for i, match in enumerate(matches):
        row_start = match.start()
        row_end = matches[i + 1].start() if i + 1 < len(matches) else len(block)
        row_text = _normalize_ws(block[row_start:row_end])
        prefix = _normalize_ws(match.group(0))
        content = row_text[len(prefix):].strip() if row_text.startswith(prefix) else row_text
        content = re.sub(r"\s*小计\s*\d*\s*$", "", content).strip()
        if content:
            rows.append({"no": match.group(1), "code": match.group(2), "content": content})
    return rows


def _normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


# ---------------------------------------------------------------------- #
# Word
# ---------------------------------------------------------------------- #
def _parse_docx_document(raw: bytes) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:  # noqa: BLE001
        raise ValueError("解析 Word 需要安装 python-docx：pip install python-docx") from exc

    try:
        doc = Document(io.BytesIO(raw))
        elements: list[ParsedElement] = []
        warnings: list[str] = []

        elements.extend(_extract_docx_container(doc, "body"))

        seen_aux_text: set[str] = set()
        for section_idx, section in enumerate(doc.sections, 1):
            for scope, container in (
                (f"header:section-{section_idx}", section.header),
                (f"footer:section-{section_idx}", section.footer),
            ):
                for element in _extract_docx_container(container, scope):
                    # Headers/footers are often linked across sections; avoid
                    # repeating identical boilerplate many times.
                    dedupe_key = f"{element.kind}:{element.text}"
                    if dedupe_key in seen_aux_text:
                        continue
                    seen_aux_text.add(dedupe_key)
                    elements.append(element)

        text = _render_elements(elements)
        if not text:
            warnings.append("Word 文档没有可读取的正文/表格文本")
        return ParsedDocument(text=text, elements=elements, warnings=warnings)
    except Exception as exc:  # noqa: BLE001
        logger.warning("DOCX parse failed: %s", exc)
        raise ValueError(f"Word 解析失败：{exc}") from exc


def _extract_docx_container(container: Any, scope: str) -> list[ParsedElement]:
    elements: list[ParsedElement] = []
    table_idx = 0
    for block in _iter_docx_blocks(container):
        block_type = block.__class__.__name__
        if block_type == "Paragraph":
            text = _normalize_multiline(block.text)
            if text:
                elements.append(ParsedElement("paragraph", text, {"scope": scope}))
        elif block_type == "Table":
            table_idx += 1
            rows = _docx_table_rows(block)
            if rows:
                elements.append(
                    ParsedElement(
                        "table",
                        _render_table(rows),
                        {"scope": scope, "table_index": table_idx},
                    )
                )
    return elements


def _iter_docx_blocks(container: Any) -> Iterable[Any]:
    """Yield paragraphs and tables in document order."""

    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    if isinstance(container, _Cell):
        parent_elm = container._tc
    elif hasattr(container, "element") and hasattr(container.element, "body"):
        parent_elm = container.element.body
    elif hasattr(container, "_element"):
        parent_elm = container._element
    else:
        return

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, container)
        elif isinstance(child, CT_Tbl):
            yield Table(child, container)


def _docx_table_rows(table: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_normalize_multiline(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(cells)
    return rows


def _normalize_multiline(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", (text or "").strip())


def _render_elements(elements: list[ParsedElement]) -> str:
    parts: list[str] = []
    for element in elements:
        scope = element.metadata.get("scope")
        table_index = element.metadata.get("table_index")
        if element.kind == "table" and scope:
            parts.append(f"### Table {table_index or ''} ({scope})\n{element.text}".strip())
        elif scope and str(scope).startswith(("header:", "footer:")):
            parts.append(f"### {scope}\n{element.text}")
        else:
            parts.append(element.text)
    return "\n\n".join(p for p in parts if p.strip())


# ---------------------------------------------------------------------- #
# Tables: CSV / Excel
# ---------------------------------------------------------------------- #
def _render_table(rows: list[list[str]]) -> str:
    """Render rows as a Markdown table."""

    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm = [r + [""] * (width - len(r)) for r in rows]
    lines = ["| " + " | ".join(_escape_md_cell(c) for c in norm[0]) + " |"]
    lines.append("| " + " | ".join("---" for _ in range(width)) + " |")
    for row in norm[1:]:
        lines.append("| " + " | ".join(_escape_md_cell(c) for c in row) + " |")
    return "\n".join(lines)


def _escape_md_cell(value: Any) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text.replace("|", r"\|").replace("\n", "<br>")


def _parse_csv_document(raw: bytes) -> ParsedDocument:
    text_in = _decode_text_bytes(raw)
    reader = csv.reader(io.StringIO(text_in))
    rows: list[list[str]] = []
    warnings: list[str] = []
    truncated = False
    for i, row in enumerate(reader):
        if i >= TABLE_MAX_ROWS:
            truncated = True
            break
        if len(row) > TABLE_MAX_COLS:
            warnings.append(f"CSV 第 {i + 1} 行超过 {TABLE_MAX_COLS} 列，已截断")
            row = row[:TABLE_MAX_COLS]
        rows.append(row)

    out = _render_table(rows)
    if truncated:
        warning = f"CSV 已截断，超出 {TABLE_MAX_ROWS} 行"
        warnings.append(warning)
        out += f"\n\n… ({warning})"
    return ParsedDocument(out, [ParsedElement("table", out, {"format": "csv"})], warnings)


def _parse_excel_document(ext: str, raw: bytes) -> ParsedDocument:
    if ext == ".xlsx":
        return _parse_xlsx_document(raw)
    return _parse_xls_document(raw)


def _parse_excel(ext: str, raw: bytes) -> str:
    """Backward-compatible helper used by older tests/imports."""

    return _parse_excel_document(ext, raw).text


def _parse_xlsx(raw: bytes) -> str:
    return _parse_xlsx_document(raw).text


def _parse_xlsx_document(raw: bytes) -> ParsedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # noqa: BLE001
        raise ValueError("解析 Excel(.xlsx) 需要安装 openpyxl：pip install openpyxl") from exc

    try:
        wb_data = load_workbook(io.BytesIO(raw), read_only=False, data_only=True)
        wb_formula = load_workbook(io.BytesIO(raw), read_only=False, data_only=False)
        try:
            parts: list[str] = []
            elements: list[ParsedElement] = []
            warnings: list[str] = []
            for ws_data, ws_formula in zip_longest(wb_data.worksheets, wb_formula.worksheets):
                if ws_data is None or ws_formula is None:
                    continue
                rows, notes = _iter_xlsx_sheet_rows(ws_data, ws_formula)
                rows = [r for r in rows if any(c.strip() for c in r)]
                if not rows and not notes:
                    continue
                sheet_parts = [f"## Sheet: {ws_data.title}"]
                if rows:
                    sheet_parts.append(_render_table(rows))
                if notes:
                    sheet_parts.append("\n".join(f"> 解析提示：{note}" for note in notes))
                    warnings.extend(f"{ws_data.title}: {note}" for note in notes)
                sheet_text = "\n".join(sheet_parts)
                parts.append(sheet_text)
                elements.append(ParsedElement("sheet", sheet_text, {"sheet": ws_data.title}))
            return ParsedDocument("\n\n".join(parts), elements, warnings)
        finally:
            wb_data.close()
            wb_formula.close()
    except Exception as exc:  # noqa: BLE001
        logger.warning("XLSX parse failed: %s", exc)
        raise ValueError(f"Excel 解析失败：{exc}") from exc


def _iter_xlsx_sheet_rows(ws_data: Any, ws_formula: Any) -> tuple[list[list[str]], list[str]]:
    max_rows = max(ws_data.max_row or 0, ws_formula.max_row or 0)
    max_cols = max(ws_data.max_column or 0, ws_formula.max_column or 0)
    row_limit = min(max_rows, TABLE_MAX_ROWS)
    col_limit = min(max_cols, TABLE_MAX_COLS)
    notes: list[str] = []

    if max_rows > TABLE_MAX_ROWS:
        notes.append(f"超过 {TABLE_MAX_ROWS} 行，已截断")
    if max_cols > TABLE_MAX_COLS:
        notes.append(f"超过 {TABLE_MAX_COLS} 列，已截断")

    merged_map = _xlsx_merged_anchor_map(ws_formula, row_limit, col_limit)
    if ws_formula.merged_cells.ranges:
        merged_preview = ", ".join(str(rng) for rng in list(ws_formula.merged_cells.ranges)[:10])
        notes.append(f"检测到合并单元格：{merged_preview}")

    hidden_rows = _preview_list(
        row for row, dim in ws_formula.row_dimensions.items() if getattr(dim, "hidden", False)
    )
    hidden_cols = _preview_list(
        col for col, dim in ws_formula.column_dimensions.items() if getattr(dim, "hidden", False)
    )
    if hidden_rows:
        notes.append(f"包含隐藏行：{hidden_rows}")
    if hidden_cols:
        notes.append(f"包含隐藏列：{hidden_cols}")

    rows: list[list[str]] = []
    for row_idx in range(1, row_limit + 1):
        cells: list[str] = []
        for col_idx in range(1, col_limit + 1):
            value = _xlsx_cell_value(ws_data, ws_formula, row_idx, col_idx)
            if (value is None or value == "") and (row_idx, col_idx) in merged_map:
                anchor_row, anchor_col = merged_map[(row_idx, col_idx)]
                value = _xlsx_cell_value(ws_data, ws_formula, anchor_row, anchor_col)

            formula_cell = ws_formula.cell(row=row_idx, column=col_idx)
            comment = getattr(formula_cell, "comment", None)
            rendered = _cell_to_str(value)
            if comment and comment.text:
                comment_text = _normalize_ws(comment.text)
                rendered = f"{rendered}（批注：{comment_text}）" if rendered else f"批注：{comment_text}"
            cells.append(rendered)
        rows.append(cells)
    return rows, notes


def _xlsx_cell_value(ws_data: Any, ws_formula: Any, row_idx: int, col_idx: int) -> Any:
    value = ws_data.cell(row=row_idx, column=col_idx).value
    if value is None:
        value = ws_formula.cell(row=row_idx, column=col_idx).value
    return value


def _xlsx_merged_anchor_map(ws: Any, row_limit: int, col_limit: int) -> dict[tuple[int, int], tuple[int, int]]:
    anchors: dict[tuple[int, int], tuple[int, int]] = {}
    for merged in ws.merged_cells.ranges:
        for row in range(merged.min_row, min(merged.max_row, row_limit) + 1):
            for col in range(merged.min_col, min(merged.max_col, col_limit) + 1):
                if (row, col) != (merged.min_row, merged.min_col):
                    anchors[(row, col)] = (merged.min_row, merged.min_col)
    return anchors


def _preview_list(values: Iterable[Any], limit: int = 12) -> str:
    items = [str(v) for _, v in zip(range(limit + 1), values)]
    if not items:
        return ""
    if len(items) > limit:
        return ", ".join(items[:limit]) + " ..."
    return ", ".join(items)


def _cell_to_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat(sep=" ", timespec="seconds")
    return str(value)


def _xls_cell_str(v: Any) -> str:
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return "" if v is None else str(v)


def _parse_xls(raw: bytes) -> str:
    return _parse_xls_document(raw).text


def _parse_xls_document(raw: bytes) -> ParsedDocument:
    try:
        import xlrd
    except ImportError as exc:  # noqa: BLE001
        raise ValueError("解析 Excel(.xls) 需要安装 xlrd：pip install xlrd") from exc

    try:
        try:
            book = xlrd.open_workbook(file_contents=raw, formatting_info=True)
        except Exception:
            book = xlrd.open_workbook(file_contents=raw)

        parts: list[str] = []
        elements: list[ParsedElement] = []
        warnings: list[str] = []
        for idx in range(book.nsheets):
            sh = book.sheet_by_index(idx)
            rows: list[list[str]] = []
            row_limit = min(sh.nrows, TABLE_MAX_ROWS)
            col_limit = min(sh.ncols, TABLE_MAX_COLS)
            notes: list[str] = []

            if sh.nrows > TABLE_MAX_ROWS:
                notes.append(f"超过 {TABLE_MAX_ROWS} 行，已截断")
            if sh.ncols > TABLE_MAX_COLS:
                notes.append(f"超过 {TABLE_MAX_COLS} 列，已截断")

            merged_map = _xls_merged_anchor_map(sh, row_limit, col_limit)
            if getattr(sh, "merged_cells", None):
                notes.append(f"检测到合并单元格：{_format_xls_merged_preview(sh.merged_cells)}")

            for i in range(row_limit):
                cells: list[str] = []
                for j in range(col_limit):
                    row, col = merged_map.get((i, j), (i, j))
                    cells.append(_xls_render_cell(book, sh, row, col))
                if any(cells):
                    rows.append(cells)

            if not rows and not notes:
                continue
            sheet_parts = [f"## Sheet: {sh.name}"]
            if rows:
                sheet_parts.append(_render_table(rows))
            if notes:
                sheet_parts.append("\n".join(f"> 解析提示：{note}" for note in notes))
                warnings.extend(f"{sh.name}: {note}" for note in notes)
            sheet_text = "\n".join(sheet_parts)
            parts.append(sheet_text)
            elements.append(ParsedElement("sheet", sheet_text, {"sheet": sh.name}))
        return ParsedDocument("\n\n".join(parts), elements, warnings)
    except Exception as exc:  # noqa: BLE001
        logger.warning("XLS parse failed: %s", exc)
        raise ValueError(f"Excel 解析失败：{exc}") from exc


def _xls_merged_anchor_map(sh: Any, row_limit: int, col_limit: int) -> dict[tuple[int, int], tuple[int, int]]:
    anchors: dict[tuple[int, int], tuple[int, int]] = {}
    for rlo, rhi, clo, chi in getattr(sh, "merged_cells", []) or []:
        for row in range(rlo, min(rhi, row_limit)):
            for col in range(clo, min(chi, col_limit)):
                if (row, col) != (rlo, clo):
                    anchors[(row, col)] = (rlo, clo)
    return anchors


def _format_xls_merged_preview(merged_cells: list[tuple[int, int, int, int]], limit: int = 10) -> str:
    preview = [f"R{rlo + 1}:R{rhi}, C{clo + 1}:C{chi}" for rlo, rhi, clo, chi in merged_cells[:limit]]
    if len(merged_cells) > limit:
        preview.append("...")
    return ", ".join(preview)


def _xls_render_cell(book: Any, sh: Any, row: int, col: int) -> str:
    try:
        import xlrd
    except ImportError:
        xlrd = None  # type: ignore[assignment]

    cell = sh.cell(row, col)
    if xlrd is not None and cell.ctype == xlrd.XL_CELL_DATE:
        try:
            return xlrd.xldate_as_datetime(cell.value, book.datemode).isoformat(sep=" ", timespec="seconds")
        except Exception:  # noqa: BLE001
            return _xls_cell_str(cell.value)
    return _xls_cell_str(cell.value)
